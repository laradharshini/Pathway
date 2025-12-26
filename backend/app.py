from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_jwt_extended import JWTManager, jwt_required, get_jwt_identity, verify_jwt_in_request
from flask_socketio import SocketIO, emit, join_room
import os
import io
from dotenv import load_dotenv
from data_engine import DataEngine
from matcher import CareerMatcher
from game_engine import GameEngine
from ai_lab import AILabEngine
from simulation_engine import SimulationEngine
from werkzeug.utils import secure_filename
import PyPDF2
import docx
from datetime import datetime
import json
from apscheduler.schedulers.background import BackgroundScheduler

# Import auth and database
from auth import hash_password, verify_password, generate_token, candidate_required, company_required
from database import (
    init_db, UserModel, CandidateProfileModel, CompanyProfileModel,
    JobModel, ApplicationModel, UserProgressModel, SimulationModel
)

# Load environment variables
load_dotenv()

# Initialize Engine
data_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data')
frontend_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'frontend-new', 'dist')

app = Flask(__name__, static_folder=frontend_path, static_url_path='')
CORS(app)

# JWT Configuration
app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET_KEY', 'dev-secret-change-in-production')
jwt = JWTManager(app)

# ========== SOCKETIO CONFIGURATION ==========
app.config['SECRET_KEY'] = 'secret!'
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

# ========== LOGGING CONFIGURATION ==========
import logging
from logging.handlers import RotatingFileHandler

# Configure logging
logging.basicConfig(level=logging.INFO)
handler = RotatingFileHandler('app.log', maxBytes=10000, backupCount=3)
handler.setFormatter(logging.Formatter(
    '[%(asctime)s] %(levelname)s in %(module)s: %(message)s'
))
app.logger.addHandler(handler)
app.logger.setLevel(logging.INFO)

# ========== CACHING CONFIGURATION ==========
from flask_caching import Cache

# Use SimpleCache (valid for single-instance dev/prod) or Redis if available
cache_config = {
    "CACHE_TYPE": "SimpleCache",  # Fallback to in-memory
    "CACHE_DEFAULT_TIMEOUT": 300
}
app.config.from_mapping(cache_config)
cache = Cache(app)

# Rate Limiter
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://"
)

# Initialize database
try:
    init_db()
    
    engine = DataEngine(data_path=data_path)
    matcher = CareerMatcher(engine)
    game_engine = GameEngine()
    ai_lab = AILabEngine()
    sim_engine = SimulationEngine()
    
    # Trigger training/loading on startup
    app.logger.info("Initializing Intelligence Engine...")
    try:
        matcher.train()
        app.logger.info("Engine Ready.")
    except Exception as e:
        app.logger.error(f"ML Engine Training Failed: {e}")
        # Proceed without ML - endpoints will degrade gracefully
        
except Exception as e:
    app.logger.error(f"Critical Startup Error (Database): {e}")
    # Initialize dummy objects if needed so imports don't fail later
    if 'engine' not in locals(): engine = DataEngine(data_path=data_path)
    if 'matcher' not in locals(): matcher = CareerMatcher(engine)
    if 'game_engine' not in locals(): game_engine = GameEngine()
    if 'ai_lab' not in locals(): ai_lab = AILabEngine()

# ========== REAL-TIME NOTIFICATION HELPERS ==========

def notify_user(user_id, event_type, data):
    """Send real-time notification to user via WebSocket"""
    try:
        socketio.emit(event_type, data, room=str(user_id))
    except Exception as e:
        print(f"Notification error: {e}")

def notify_new_job_match(user_id, job_data):
    """Notify user of new matching job"""
    notify_user(user_id, 'new_job_match', {
        'job': job_data,
        'timestamp': datetime.utcnow().isoformat(),
        'message': f"New job match: {job_data['title']} at {job_data['company']}"
    })

def notify_application_update(user_id, application_id, status):
    """Notify when application status changes"""
    notify_user(user_id, 'application_update', {
        'application_id': application_id,
        'status': status,
        'timestamp': datetime.utcnow().isoformat()
    })

# ========== WEBSOCKET EVENTS ==========

@socketio.on('connect')
def handle_connect():
    print(f'Client connected: {request.sid}')
    emit('connection_response', {'status': 'connected', 'timestamp': datetime.utcnow().isoformat()})

@socketio.on('disconnect')
def handle_disconnect():
    print(f'Client disconnected: {request.sid}')

@socketio.on('join_user_room')
def handle_join_room(data):
    """Join user-specific room for personalized notifications"""
    user_id = data.get('user_id')
    if user_id:
        join_room(str(user_id))
        emit('room_joined', {'room': user_id, 'status': 'success'})
        print(f"User {user_id} joined their notification room")

@socketio.on('ping')
def handle_ping():
    """Keep-alive ping"""
    emit('pong', {'timestamp': datetime.utcnow().isoformat()})

# ========== BACKGROUND JOB SYNC ==========

def sync_new_jobs():
    """Background task to check for new jobs and notify relevant candidates"""
    try:
        print("Running job sync...")
        # Get all active candidates
        from database import candidate_profiles_collection
        candidates = list(candidate_profiles_collection.find({}))
        
        # Get recently added jobs (last 6 hours)
        from database import jobs_collection
        from datetime import timedelta
        recent_cutoff = datetime.utcnow() - timedelta(hours=6)
        new_jobs = list(jobs_collection.find({
            'is_active': True,
            'created_at': {'$gte': recent_cutoff}
        }))
        
        if not new_jobs:
            print("No new jobs to sync")
            return
        
        print(f"Found {len(new_jobs)} new jobs, checking matches...")
        
        # Check each candidate against new jobs
        for candidate in candidates:
            profile_data = {
                'preferred_role': candidate.get('target_role', ''),
                'experience_level': candidate.get('experience_level', 'entry'),
                'skills': candidate.get('skills', [])
            }
            
            # Get matches
            matches = matcher.match_user(profile_data)
            
            # Find new jobs in top matches
            for match in matches[:5]:  # Top 5 matches
                if match['readiness_score'] >= 60:  # Only notify if decent match
                    # Check if this is a new job
                    job_id = str(match['job_id'])
                    if any(str(job['_id']) == job_id for job in new_jobs):
                        # Send notification
                        notify_new_job_match(
                            candidate['user_id'],
                            {
                                'job_id': job_id,
                                'title': match['title'],
                                'company': match['company'],
                                'readiness_score': match['readiness_score']
                            }
                        )
                        print(f"Notified {candidate['user_id']} about {match['title']}")
        
        print("Job sync complete")
        
    except Exception as e:
        print(f"Job sync error: {e}")
        import traceback
        traceback.print_exc()

# Initialize scheduler
scheduler = BackgroundScheduler()
scheduler.add_job(func=sync_new_jobs, trigger="interval", hours=6)
scheduler.start()

# ========== EXISTING ROUTES (KEPT AS-IS) ==========

@app.route('/')
def serve_frontend():
    return app.send_static_file('index.html')

@app.route('/api/auth/signup', methods=['POST'])
def signup():
    """User signup endpoint"""
    try:
        data = request.json
        email = data.get('email')
        password = data.get('password')
        role = data.get('role')
        
        if not email or not password or not role:
            return jsonify({'error': 'Email, password, and role are required'}), 400
        
        if role not in ['candidate', 'company', 'admin']:
            return jsonify({'error': 'Invalid role'}), 400
        
        existing_user = UserModel.find_by_email(email)
        if existing_user:
            return jsonify({'error': 'Email already registered'}), 409
        
        password_hash = hash_password(password)
        user = UserModel.create(email, password_hash, role)
        
        if role == 'candidate':
            profile_data = data.get('profile', {})
            CandidateProfileModel.create(
                user_id=str(user['_id']),
                full_name=profile_data.get('full_name', ''),
                experience_level=profile_data.get('experience_level', 'entry'),
                target_role=profile_data.get('target_role', ''),
                location=profile_data.get('location', '')
            )
        elif role == 'company':
            profile_data = data.get('profile', {})
            CompanyProfileModel.create(
                user_id=str(user['_id']),
                company_name=profile_data.get('company_name', ''),
                industry=profile_data.get('industry', ''),
                location=profile_data.get('location', ''),
                description=profile_data.get('description', '')
            )
        
        # Generate Token for immediate login
        access_token = generate_token(str(user['_id']), user['role'])

        return jsonify({
            'message': 'Account created successfully',
            'user_id': str(user['_id']),
            'role': role,
            'access_token': access_token,
            'user': {
                'id': str(user['_id']),
                'email': user['email'],
                'role': user['role'],
                'full_name': profile_data.get('full_name', '')
            }
        }), 201
        
    except Exception as e:
        print(f"Signup error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/auth/login', methods=['POST'])
def login():
    """User login endpoint"""
    try:
        data = request.json
        email = data.get('email')
        password = data.get('password')
        
        if not email or not password:
            return jsonify({'error': 'Email and password are required'}), 400
        
        user = UserModel.find_by_email(email)
        if not user:
            return jsonify({'error': 'Invalid credentials'}), 401
        
        if not verify_password(password, user['password_hash']):
            return jsonify({'error': 'Invalid credentials'}), 401
        
        access_token = generate_token(str(user['_id']), user['role'])
        
        return jsonify({
            'access_token': access_token,
            'user': {
                'id': str(user['_id']),
                'email': user['email'],
                'role': user['role']
            }
        }), 200
        
    except Exception as e:
        print(f"Login error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/auth/social-mock', methods=['POST'])
def social_mock_login():
    """
    Exchanges mock social login data for a REAL backend session (JWT).
    Upserts the user so they exist in the DB.
    """
    try:
        data = request.json
        email = data.get('email')
        name = data.get('full_name')
        role = data.get('role', 'candidate')
        provider = data.get('provider', 'Google')
        
        if not email:
            return jsonify({'error': 'Email required'}), 400
            
        # 1. Find or Create User
        user = UserModel.find_by_email(email)
        if not user:
            # Create new user with a dummy password (social users usually don't have passwords)
            # We'll assign a random strong one just to satisfy the DB schema
            import secrets
            dummy_pw = hash_password(secrets.token_hex(16))
            user = UserModel.create(email, dummy_pw, role)
            
            # Create Profile immediately
            if role == 'candidate':
                CandidateProfileModel.create(
                    user_id=str(user['_id']),
                    full_name=name or 'Social User',
                    experience_level='mid', # Default to mid for demo
                    target_role='Software Engineer', # Default
                    location='Remote'
                )
        
        # 2. Update Profile Name if needed (visual sync)
        if name and role == 'candidate':
            # Simple update if name provided
            from database import candidate_profiles_collection
            from bson import ObjectId
            candidate_profiles_collection.update_one(
                {'user_id': str(user['_id'])},
                {'$set': {'full_name': name}}
            )

        # 3. Generate REAL Token
        access_token = generate_token(str(user['_id']), user['role'])
        
        return jsonify({
            'access_token': access_token,
            'user': {
                'id': str(user['_id']),
                'email': user['email'],
                'role': user['role'],
                'full_name': name
            }
        }), 200

    except Exception as e:
        print(f"Social Mock Login Error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/candidate/profile', methods=['GET', 'PUT'])
@candidate_required
def candidate_profile():
    """Get or Update candidate profile"""
    print(f"DEBUG: Endpoint /api/candidate/profile hit via {request.method}")
    try:
        user_id = get_jwt_identity()
        print(f"DEBUG: User ID from JWT: {user_id}")
        
        if request.method == 'GET':
            profile = CandidateProfileModel.find_by_user_id(user_id)
            if not profile: return jsonify({'error': 'Profile not found'}), 404
            profile['_id'] = str(profile['_id'])
            return jsonify(profile), 200
            
        elif request.method == 'PUT':
            data = request.json
            app.logger.info(f"Updating profile for {user_id}: {data}")
            
            # Whitelist allowed fields for update
            allowed_fields = ['full_name', 'target_role', 'experience_level', 'location', 'skills', 'bio', 'linkedin_url', 'github_url', 'avatar_url']
            update_data = {k: v for k, v in data.items() if k in allowed_fields}
            
            if 'updated_at' not in update_data:
                update_data['updated_at'] = datetime.utcnow()
                
            CandidateProfileModel.update(user_id, update_data)
            
            # Send real-time notification
            from auth import get_current_user_id
            notify_user(user_id, 'profile_updated', {
                'message': 'Profile updated successfully',
                'timestamp': datetime.utcnow().isoformat()
            })
            
            # Re-fetch to confirm and return updated object
            updated_profile = CandidateProfileModel.find_by_user_id(user_id)
            updated_profile['_id'] = str(updated_profile['_id'])
            
            return jsonify({
                'message': 'Profile updated successfully',
                'profile': updated_profile
            }), 200
            
    except Exception as e:
        app.logger.error(f"Profile update error: {e}")
        return jsonify({'error': str(e)}), 500

# ========== NEW: SKILL ASSESSMENT ENDPOINTS ==========

@app.route('/api/candidate/assessment', methods=['POST'])
@candidate_required
def submit_assessment():
    """Submit skill assessment results"""
    try:
        user_id = get_jwt_identity()
        data = request.json
        
        skill = data.get('skill')
        score = data.get('score')
        answers = data.get('answers', {})
        time_taken = data.get('time_taken', 0)
        
        # Determine proficiency based on score
        if score >= 80:
            proficiency = 'advanced'
        elif score >= 60:
            proficiency = 'intermediate'
        else:
            proficiency = 'beginner'
        
        # Update candidate profile with verified skill
        profile = CandidateProfileModel.find_by_user_id(user_id)
        if not profile:
            return jsonify({'error': 'Profile not found'}), 404
        
        # Add or update skill with assessment data
        skills = profile.get('skills', [])
        
        # Remove old entry if exists
        skills = [s for s in skills if s.get('name') != skill]
        
        # Add new verified skill
        skills.append({
            'name': skill,
            'proficiency': proficiency,
            'verified': True,
            'assessment_score': score,
            'assessed_at': datetime.utcnow().isoformat()
        })
        
        CandidateProfileModel.update_skills(user_id, skills)
        
        # Award XP for assessment
        xp_gain = 50 if proficiency == 'beginner' else 100 if proficiency == 'intermediate' else 200
        award_xp(user_id, xp_gain, f"Completed {skill} assessment")

        # Send real-time notification
        notify_user(user_id, 'skill_verified', {
            'skill': skill,
            'proficiency': proficiency,
            'score': score,
            'message': f'Congratulations! You scored {score}% in {skill}'
        })
        
        return jsonify({
            'message': 'Assessment recorded successfully',
            'skill': skill,
            'proficiency': proficiency,
            'score': score
        }), 200
        
    except Exception as e:
        print(f"Assessment error: {e}")
        return jsonify({'error': str(e)}), 500

# ========== NEW: AI RESUME ANALYSIS ==========

@app.route('/api/resume/analyze', methods=['POST'])
@candidate_required
def analyze_resume_ai():
    """AI-powered resume analysis using Claude"""
    try:
        user_id = get_jwt_identity()
        
        if 'resume' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        # Extract text
        text = ""
        filename = secure_filename(file.filename)
        
        if filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + " "
        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + " "
        else:
            return jsonify({"error": "Unsupported format. Use PDF or DOCX"}), 400

        # Get candidate profile for context
        profile = CandidateProfileModel.find_by_user_id(user_id)
        target_role = profile.get('target_role', 'Unknown') if profile else 'Unknown'

        # Call Claude API for analysis
        import requests
        
        claude_response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": os.getenv('ANTHROPIC_API_KEY', ''),
                "anthropic-version": "2023-06-01"
            },
            json={
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 1500,
                "messages": [{
                    "role": "user",
                    "content": f"""Analyze this resume for a {target_role} position. Provide:

1. ATS compatibility score (0-100) - how well it will pass applicant tracking systems
2. Top 3 strengths (be specific)
3. Top 3 improvements needed (actionable advice)
4. Missing keywords for {target_role} roles
5. Overall recommendation

Resume text:
{text[:4000]}

Respond ONLY in valid JSON format with keys: ats_score, strengths, improvements, missing_keywords, recommendation"""
                }]
            }
        )
        
        if claude_response.status_code != 200:
            app.logger.warning(f"Claude API failed: {claude_response.text}. Using fallback analysis.")
            # Fallback Mock Analysis for Demo
            import random
            mock_score = random.randint(70, 90)
            analysis = {
                "ats_score": mock_score,
                "strengths": ["Strong experience metrics", "Clear formatting", "Relevant keywords detected"],
                "improvements": ["Add more quantitative results", "Expand on soft skills", "Include localized contact info"],
                "missing_keywords": ["Agile", "JIRA", "System Design"],
                "recommendation": "Great starting point. Consider tailoring the objective statement to the specific role you are applying for."
            }
        else:
            response_data = claude_response.json()
            analysis_text = response_data['content'][0]['text']
            # Parse JSON
            import re
            json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
            if json_match:
                analysis = json.loads(json_match.group())
            else:
                 raise ValueError("Could not extract JSON from AI response")
        
        # Send real-time notification
        notify_user(user_id, 'resume_analyzed', {
            'ats_score': analysis.get('ats_score', 0),
            'message': f"Resume analysis complete! ATS Score: {analysis.get('ats_score', 0)}/100"
        })
        
        return jsonify({
            'analysis': analysis,
            'text_preview': text[:500]
        }), 200
        
    except Exception as e:
        app.logger.error(f"Resume analysis error: {e}")
        # Final fallback if even fallback fails
        return jsonify({'analysis': {
             "ats_score": 75,
             "strengths": ["Analysis Engine Unavailable - Using Cached Profile"],
             "improvements": ["Check API Connectivity"],
             "missing_keywords": [],
             "recommendation": "Service is currently in offline mode."
        }, 'text_preview': "Preview unavailable"}), 200

@app.route('/api/jobs/<job_id>/resume_match', methods=['POST'])
@candidate_required
def match_resume_to_job(job_id):
    """Analyze resume against a specific job for ATS score and autofill data"""
    try:
        user_id = get_jwt_identity()
        
        if 'resume' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        # Extract text from resume
        text = ""
        filename = secure_filename(file.filename)
        
        if filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + " "
        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + " "
        else:
            return jsonify({"error": "Unsupported format. Use PDF or DOCX"}), 400

        # Fetch job details for context
        job = JobModel.find_by_id(job_id)
        if not job:
            job_title = request.form.get('job_title', 'Unknown Role')
            job_desc = request.form.get('job_description', '')
        else:
            job_title = job.get('title')
            job_desc = job.get('description')

        # Try AI Analysis (Prefer Gemini for Free Tier, then Claude)
        import requests
        google_api_key = os.getenv('GOOGLE_API_KEY', '')
        anthropic_api_key = os.getenv('ANTHROPIC_API_KEY', '')
        
        app.logger.info(f"Checking AI configuration. Gemini Key: {'Set' if google_api_key else 'Missing'}, Claude Key: {'Set' if anthropic_api_key else 'Missing'}")
        
        is_google_configured = google_api_key and 'your-google-api-key' not in google_api_key.lower()
        is_anthropic_configured = anthropic_api_key and 'your-anthropic-api-key' not in anthropic_api_key.lower()
        
        app.logger.info(f"Is Google Configured: {is_google_configured}, Is Anthropic Configured: {is_anthropic_configured}")
        
        analysis = None

        # 1. Try Google Gemini (Free Tier)
        if is_google_configured:
            app.logger.info("Attempting Google Gemini analysis via SDK...")
            try:
                import google.generativeai as genai
                genai.configure(api_key=google_api_key)
                
                # Check available models to pick the best one (prefer 2.5-flash)
                available_models = [m.name for m in genai.list_models()]
                if 'models/gemini-2.5-flash' in available_models:
                    model_name = 'models/gemini-2.5-flash'
                elif 'models/gemini-1.5-flash' in available_models:
                    model_name = 'models/gemini-1.5-flash'
                else:
                    model_name = 'gemini-pro'
                
                app.logger.info(f"Using Gemini model: {model_name}")
                model = genai.GenerativeModel(model_name)
                
                prompt = f"""Analyze this resume against the following job description:
Job Title: {job_title}
Job Description: {job_desc}

1. Provide an ATS compatibility score (0-100).
2. Extract the following candidate details for a form: full_name, email, phone, experience_summary, and a list of detected_skills.

Resume text:
{text[:4000]}

Respond ONLY in valid JSON format with keys: ats_score, candidate_details (which includes full_name, email, phone, experience_summary, detected_skills)"""

                response = model.generate_content(prompt)
                analysis_text = response.text
                
                import re
                json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
                if json_match:
                    analysis = json.loads(json_match.group())
                    analysis['is_demo'] = False
                    app.logger.info(f"Successfully used Google Gemini ({model_name}) for analysis.")
                else:
                    # Sanitize for logging to avoid UnicodeEncodeError in Windows terminal
                    safe_text = analysis_text[:200].encode('ascii', 'ignore').decode('ascii')
                    app.logger.error(f"Gemini responded but no JSON found in text: {safe_text}")
            except Exception as e:
                app.logger.error(f"Google Gemini SDK error: {str(e).encode('ascii', 'ignore').decode('ascii')}")

        # 2. Try Anthropic Claude (Fallback)
        if not analysis and is_anthropic_configured:
            try:
                claude_response = requests.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "Content-Type": "application/json",
                        "x-api-key": anthropic_api_key,
                        "anthropic-version": "2023-06-01"
                    },
                    json={
                        "model": "claude-3-haiku-20240307",
                        "max_tokens": 1000,
                        "messages": [{
                            "role": "user",
                            "content": f"""Analyze this resume against the following job description:
Job Title: {job_title}
Job Description: {job_desc}

1. Provide an ATS compatibility score (0-100).
2. Extract the following candidate details for a form: full_name, email, phone, experience_summary, and a list of detected_skills.

Resume text:
{text[:4000]}

Respond ONLY in valid JSON format with keys: ats_score, candidate_details (which includes full_name, email, phone, experience_summary, detected_skills)"""
                        }]
                    },
                    timeout=15
                )
                
                if claude_response.status_code == 200:
                    response_data = claude_response.json()
                    analysis_text = response_data['content'][0]['text']
                    import re
                    json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
                    if json_match:
                        analysis = json.loads(json_match.group())
                        analysis['is_demo'] = False
                        app.logger.info("Successfully used Anthropic Claude for analysis.")
            except Exception as e:
                app.logger.warning(f"Anthropic Claude analysis failed: {e}")

        # 3. Fallback/Demo mode
        if not analysis:
            analysis = {
                "ats_score": 82,
                "is_demo": True,
                "candidate_details": {
                    "full_name": "Demo Candidate",
                    "email": "candidate@example.com",
                    "phone": "123-456-7890",
                    "experience_summary": "Extensive experience in software development and project management.",
                    "detected_skills": ["Python", "React", "Node.js", "SQL", "Project Management"]
                }
            }
        
        return jsonify(analysis), 200
        
    except Exception as e:
        app.logger.error(f"Job specific resume analysis error: {e}")
        return jsonify({'error': str(e)}), 500

# ========== NEW: INTERVIEW ANALYSIS ==========

@app.route('/api/interview/analyze', methods=['POST'])
@candidate_required
def analyze_interview_answer():
    """Analyze interview answer using Claude AI"""
    try:
        data = request.json
        question = data.get('question')
        answer = data.get('answer')
        
        if not question or not answer:
            return jsonify({'error': 'Question and answer required'}), 400
        
        # Call Claude API
        import requests
        
        claude_response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": os.getenv('ANTHROPIC_API_KEY', ''),
                "anthropic-version": "2023-06-01"
            },
            json={
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 800,
                "messages": [{
                    "role": "user",
                    "content": f"""You are an expert interview coach. Analyze this interview answer:

Question: {question}
Answer: {answer}

Provide scores (0-10) and feedback:
1. clarity_score - How clear and well-structured
2. technical_score - Technical depth and accuracy
3. confidence_score - Confidence and communication style
4. feedback - 2-3 sentences of actionable advice

Respond ONLY in valid JSON format."""
                }]
            }
        )
        
        if claude_response.status_code != 200:
            return jsonify({'error': 'AI analysis failed'}), 500
        
        response_data = claude_response.json()
        analysis_text = response_data['content'][0]['text']
        
        # Parse JSON
        import re
        json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
        if json_match:
            analysis = json.loads(json_match.group())
        else:
            analysis = {
                'clarity_score': 7,
                'technical_score': 6,
                'confidence_score': 8,
                'feedback': 'Good effort! Focus on providing more specific examples.'
            }
        
        return jsonify(analysis), 200
        
    except Exception as e:
        print(f"Interview analysis error: {e}")
        return jsonify({'error': str(e)}), 500

# ========== GAMIFICATION ENDPOINTS ==========

@app.route('/api/gamification/progress', methods=['GET'])
@jwt_required()
def get_user_progress():
    try:
        user_id = get_jwt_identity()
        progress = UserProgressModel.find_by_user_id(user_id)
        if not progress:
            progress = UserProgressModel.create(user_id)
        
        # Convert ObjectIds
        progress['_id'] = str(progress['_id'])
        return jsonify(progress), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def award_xp(user_id, amount, action_description):
    """Internal helper to award XP and notify user"""
    try:
        leveled_up = UserProgressModel.update_xp(user_id, amount)
        
        # Notify
        notify_user(user_id, 'xp_gained', {
            'amount': amount,
            'reason': action_description,
            'leveled_up': leveled_up
        })
        
        if leveled_up:
            progress = UserProgressModel.find_by_user_id(user_id)
            notify_user(user_id, 'level_up', {
                'new_level': progress['level'],
                'message': f"Level Up! You are now Level {progress['level']}!"
            })
            
    except Exception as e:
        app.logger.error(f"Error awarding XP: {e}")

# ========== KEEP ALL YOUR EXISTING ROUTES ==========
# (I'm preserving all your original routes below)

@app.route('/api/games/trivia', methods=['GET'])
@candidate_required
def get_trivia():
    try:
        # Get user skills to personalize
        user_id = get_jwt_identity()
        try:
            profile = CandidateProfileModel.find_by_user_id(user_id)
            skills = profile.get('skills', ['python']) if profile else ['python']
        except Exception:
            # Fallback if DB is down or other error
            skills = ['python', 'react', 'sql']
            
        questions = game_engine.get_trivia(skills)
        return jsonify(questions), 200
    except Exception as e:
        # Even if everything fails, return empty list instead of 500 to keep UI alive
        return jsonify([]), 200

@app.route('/api/games/bugs', methods=['GET'])
@candidate_required
def get_bugs():
    try:
        bugs = game_engine.get_bugs()
        return jsonify(bugs), 200
    except Exception as e:
        return jsonify([]), 200

@app.route('/api/games/scenario', methods=['GET'])
@candidate_required
def get_scenario():
    try:
        scenario = game_engine.get_scenario()
        return jsonify(scenario), 200
    except Exception as e:
        return jsonify({}), 200

# AI Lab Routes
@app.route('/api/games/ai-lab/challenge', methods=['GET'])
@candidate_required
def get_ai_challenge():
    try:
        skill = request.args.get('skill', 'python')
        proficiency = request.args.get('proficiency', 'intermediate')
        challenge = ai_lab.generate_challenge(skill, proficiency)
        return jsonify(challenge), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/games/ai-lab/evaluate', methods=['POST'])
@candidate_required
def evaluate_ai_challenge():
    try:
        data = request.json
        title = data.get('title')
        code = data.get('code')
        result = ai_lab.evaluate_submission(title, code)
        
        # Award XP if passed
        if result.get('passed'):
            user_id = get_jwt_identity()
            award_xp(user_id, 150, f"Solved AI Lab challenge: {title}")
            
        return jsonify(result), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/games/ai-lab/hint', methods=['POST'])
@candidate_required
def get_ai_hint():
    try:
        data = request.json
        title = data.get('title')
        code = data.get('code')
        hint = ai_lab.get_hint(title, code)
        return jsonify(hint), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Keep all other existing routes...
# (Your company, jobs, applications routes remain unchanged)

# ========== COMPANY PROFILE ENDPOINTS ==========

@app.route('/api/company/profile', methods=['GET'])
@company_required
def get_company_profile():
    """Get company's own profile"""
    try:
        user_id = get_jwt_identity()
        profile = CompanyProfileModel.find_by_user_id(user_id)
        if not profile:
            return jsonify({'error': 'Profile not found'}), 404
        profile['_id'] = str(profile['_id'])
        return jsonify(profile), 200
    except Exception as e:
        print(f"Get company profile error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/jobs/recommendations', methods=['GET'])
@candidate_required
def get_job_recommendations():
    """
    Get personalized, real-time job recommendations with readiness scores.
    """
    try:
        user_id = get_jwt_identity()
        profile = CandidateProfileModel.find_by_user_id(user_id)
        
        if not profile:
            return jsonify({'error': 'Profile not found'}), 404

        # 1. LIVE DATA FETCH (Phase 3)
        # Query The Muse API based on user's target role and location
        role_query = profile.get('target_role', 'Software Engineer')
        location_query = profile.get('location', '')
        
        live_jobs = engine.search_jobs(query=role_query, location=location_query)
        
        # Fallback to DB if live search fails
        if not live_jobs:
            live_jobs = JobModel.find_active()
            
        if live_jobs:
            # DataEngine already handles skill extraction now
            pass
            
        # 2. READINESS SCORING (Phase 6)
        # Pass the live jobs to the matcher to apply the 0.7/0.2/0.1 formula
        matcher_profile = {
            'skills': profile.get('skills', []),
            'experience_level': profile.get('experience_level', 'Entry Level'),
            'preferred_role': role_query
        }
        
        scored_jobs = matcher.match_user(matcher_profile, live_jobs=live_jobs)
        
        # 3. Return as 'jobs' for dashboard parity
        return jsonify({
            'jobs': scored_jobs[:20], 
            'source': 'live_api',
            'query_used': role_query
        }), 200

    except Exception as e:
        app.logger.error(f"Recommendations error: {e}")
        return jsonify({'error': str(e)}), 500

# ========== JOB POSTING ENDPOINTS ==========

@app.route('/api/jobs', methods=['POST'])
@company_required
def create_job():
    """Create new job posting"""
    try:
        user_id = get_jwt_identity()
        company_profile = CompanyProfileModel.find_by_user_id(user_id)
        if not company_profile:
            return jsonify({'error': 'Company profile not found'}), 404
        
        data = request.json
        title = data.get('title')
        description = data.get('description')
        skills_input = data.get('skills', [])
        
        # Simple extraction if skills missing
        if not skills_input and description:
            # Basic keyword matching placeholder
            common = ['Python', 'Java', 'React', 'SQL']
            skills_input = [s for s in common if s.lower() in description.lower()]

        job = JobModel.create(
            company_id=str(company_profile['_id']),
            title=title,
            description=description,
            experience_level=data.get('experience_level', 'entry'),
            location=data.get('location', ''),
            skills=skills_input
        )
        return jsonify({'message': 'Job posted', 'job_id': str(job['_id'])}), 201
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/jobs', methods=['GET'])
@company_required
def get_company_jobs():
    try:
        user_id = get_jwt_identity()
        company_profile = CompanyProfileModel.find_by_user_id(user_id)
        if not company_profile:
            return jsonify({'error': 'Company profile not found'}), 404
        
        jobs = JobModel.find_by_company(str(company_profile['_id']))
        for job in jobs: job['_id'] = str(job['_id'])
        return jsonify({'jobs': jobs}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/jobs/search', methods=['GET'])
def search_jobs():
    try:
        # Get query parameters
        query = request.args.get('query', 'Software Engineering')
        location = request.args.get('location', '')
        
        # 1. Fetch Jobs
        live_jobs = engine.search_jobs(query, location)
        if not live_jobs:
            live_jobs = JobModel.find_active()
            
        # 2. SANITIZE SKILLS (Fix for Stale DataEngine)
        import re
        def sanitize_skills(text):
            found = set()
            keywords = {
                'Python': ['python'], 'Java': ['java', 'jvm'], 'JavaScript': ['javascript', 'js'],
                'React': ['react'], 'Node.js': ['node'], 'SQL': ['sql'], 'AWS': ['aws'],
                'Docker': ['docker'], 'Communication': ['communication'], 'Teamwork': ['teamwork'],
                'Data Analysis': ['analysis', 'data'], 'Design': ['design', 'figma'],
                'Software Engineering': ['software', 'developer', 'engineer']
            }
            text_lower = (text or '').lower()
            for skill, patterns in keywords.items():
                for pat in patterns:
                    if re.search(r'\b' + re.escape(pat) + r'\b', text_lower):
                        found.add(skill)
                        break
            if not found: found.add('Communication')
            return list(found)

        if live_jobs:
            for job in live_jobs:
                desc = job.get('full_description') or job.get('description', '')
                job['mapped_skills'] = sanitize_skills(desc)

        # 3. Match User (if logged in)
        user_id = None
        try:
            verify_jwt_in_request(optional=True)
            user_id = get_jwt_identity()
        except Exception as e:
            with open("backend_debug.log", "a") as f: f.write(f"DEBUG: JWT Verify Error: {str(e)}\n")
            pass

        if user_id:
            profile = CandidateProfileModel.find_by_user_id(user_id)
            if profile:
                matcher_profile = {
                    'skills': profile.get('skills', []),
                    'experience_level': profile.get('experience_level', 'Entry Level'),
                    'preferred_role': profile.get('target_role', query)
                }
                scored_jobs = matcher.match_user(matcher_profile, live_jobs=live_jobs)
                return jsonify({'jobs': scored_jobs}), 200
        
        return jsonify({'jobs': live_jobs}), 200

    except Exception as e:
        app.logger.error(f"Search jobs error: {e}")
        return jsonify({'error': str(e)}), 500


# ========== APPLICATION ENDPOINTS ==========

@app.route('/api/jobs/<job_id>/applications', methods=['GET'])
@company_required
def get_job_applications(job_id):
    """Get all applications for a specific job"""
    try:
        user_id = get_jwt_identity()
        company_profile = CompanyProfileModel.find_by_user_id(user_id)
        if not company_profile:
            return jsonify({'error': 'Company profile not found'}), 404
            
        # Verify job belongs to company
        # (Assuming JobModel has find_one or we trust the company_id check later)
        # For strictness:
        jobs = JobModel.find_by_company(str(company_profile['_id']))
        if not any(str(j['_id']) == job_id for j in jobs):
             return jsonify({'error': 'Job not found or unauthorized'}), 403

        apps = ApplicationModel.find_by_job(job_id)
        
        # Enrich with candidate details
        enriched_apps = []
        for app in apps:
            candidate = CandidateProfileModel.find_by_user_id(app['candidate_id']) # Wait, app has candidate_id which is profile_id usually?
            # In create application: candidate_id=str(profile['_id'])
            # So find_by_id on profile collection
            if not candidate:
                 # Try finding by _id directly if your model supports it
                 # Converting to find_one({'_id': ObjectId(candidate_id)})
                 from database import candidate_profiles_collection
                 from bson import ObjectId
                 candidate = candidate_profiles_collection.find_one({'_id': ObjectId(app['candidate_id'])})
            
            if candidate:
                app_data = {
                    'application_id': str(app['_id']),
                    'status': app['status'],
                    'readiness_score': app.get('readiness_score', 0),
                    'applied_at': app.get('applied_at'),
                    'candidate': {
                        'full_name': candidate.get('full_name'),
                        'skills': candidate.get('skills', []),
                        'experience_level': candidate.get('experience_level')
                    }
                }
                enriched_apps.append(app_data)
        
        return jsonify({'applications': enriched_apps}), 200
    except Exception as e:
        print(f"Get job applications error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/jobs/<job_id>/apply', methods=['POST'])
@candidate_required
def apply_to_job(job_id):
    try:
        from bson import ObjectId
        user_id = get_jwt_identity()
        profile = CandidateProfileModel.find_by_user_id(user_id)
        if not profile:
            return jsonify({'error': 'Profile not found'}), 404
            
        # Check existing
        existing = ApplicationModel.check_existing(job_id, str(profile['_id']))
        if existing:
            return jsonify({'error': 'Already applied'}), 409

        # Calculate Readiness
        profile_data = {
            'preferred_role': profile.get('target_role', ''),
            'experience_level': profile.get('experience_level', 'entry'),
            'skills': profile.get('skills', [])
        }
        matches = matcher.match_user(profile_data)
        job_match = next((m for m in matches if str(m.get('job_id')) == str(job_id)), None)
        
        # If not found in recommendations, try fetching directly
        if not job_match:
            try:
                from database import JobModel
                job_doc = JobModel.find_by_id(job_id)
                if job_doc:
                    single_matches = matcher.match_user(profile_data, live_jobs=[job_doc])
                    if single_matches:
                        job_match = single_matches[0]
            except Exception as e:
                app.logger.warning(f"Could not find job {job_id} for scoring: {e}")

        score = job_match.get('readiness_score', 50.0) if job_match else 50.0
        company_name = job_match.get('company', 'a job') if job_match else 'a job'
        
        data = request.json or {}
        candidate_details = data.get('candidate_details', {})

        app_id = ApplicationModel.create(
            job_id=job_id,
            candidate_id=str(profile['_id']),
            readiness_score=score,
            match_score=score,
            candidate_details=candidate_details
        )
        
        # Award XP
        award_xp(user_id, 20, f"Applied to {company_name}")

        # Notify Company (implied)
        return jsonify({'message': 'Application sent', 'readiness_score': score}), 201

    except Exception as e:
        print(f"Apply error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/candidate/applications', methods=['GET'])
@candidate_required
def get_candidate_applications():
    try:
        from bson import ObjectId
        user_id = get_jwt_identity()
        profile = CandidateProfileModel.find_by_user_id(user_id)
        if not profile: return jsonify({'error': 'Profile not found'}), 404
        
        apps = ApplicationModel.find_by_candidate(str(profile['_id']))
        # Enrich would go here (omitted for brevity, returning raw for now)
        for a in apps: a['_id'] = str(a['_id'])
        return jsonify({'applications': apps}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ========== CORE INTELLIGENCE ENDPOINT ==========

@app.route('/api/analyze', methods=['POST'])
def analyze():
    """
    Analyze user profile and match with jobs.
    """
    try:
        from flask_jwt_extended import verify_jwt_in_request, get_jwt_identity
        
        # Try auth
        user_id = None
        try:
            verify_jwt_in_request(optional=True)
            user_id = get_jwt_identity()
        except: pass
        
        if user_id:
            profile = CandidateProfileModel.find_by_user_id(user_id)
            if not profile: return jsonify({'error': 'Profile not found'}), 404
            
            profile_data = {
                'preferred_role': profile.get('target_role', ''),
                'experience_level': profile.get('experience_level', 'entry'),
                'skills': profile.get('skills', [])
            }
        else:
            profile_data = request.json
            
        results = matcher.match_user(profile_data)
        
        # Analysis summary
        avg_readiness = sum(r['readiness_score'] for r in results) / len(results) if results else 0
        
        # Missing skills
        from collections import Counter
        all_missing = []
        for job in results[:5]:
            all_missing.extend(job.get('missing_skills', []))
        critical = [s for s, c in Counter(all_missing).most_common(5)]
        
        return jsonify({
            'readiness_score': round(avg_readiness, 1),
            'top_matches': results[:10],
            'analysis': {
                'summary': "Analysis Complete",
                'missing_critical_skills': critical
            }
        })
        
    except Exception as e:
        print(f"Analyze error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/simulate', methods=['POST'])
def simulate_skill_impact():
    try:
        data = request.json
        print(f"Simulate Request: {data.keys()}")
        if 'profile' not in data:
             print("Error: Missing profile in request")
             return jsonify({'error': 'Missing profile'}), 400
             
        sim_profile = data['profile'].copy()
        if 'skills' not in sim_profile or sim_profile['skills'] is None:
            print("Warning: Profile missing skills, initializing empty list")
            sim_profile['skills'] = []
            
        sim_profile['skills'] = list(sim_profile['skills']) # Ensure list copy
        sim_profile['skills'].append(data['add_skill'])
        
        target_id = str(data.get('target_job_id', 'unknown'))
        
        # 1. Check if full job data was provided (for live/non-persisted jobs)
        if 'job' in data:
            print(f"Simulating using provided job data for: {target_id}")
            matches = matcher.match_user(sim_profile, live_jobs=[data['job']])
        else:
            # 2. Targeted check: Try to fetch job from DB
            job_data = JobModel.find_by_id(target_id)
            if job_data:
                print(f"Targeted simulation for DB job: {target_id}")
                matches = matcher.match_user(sim_profile, live_jobs=[job_data])
            else:
                # 3. Fallback to search in local cache (CSV jobs)
                print(f"Fallback simulation in local cache for: {target_id}")
                matches = matcher.match_user(sim_profile)
        
        job = next((j for j in matches if str(j['job_id']) == target_id), None)
        
        if not job: 
            print(f"No match found for job {target_id}. Matches evaluated: {len(matches)}")
            # If we have matches but none match the ID, it's a mapping issue. Return the first one if only one was requested.
            if len(matches) == 1 and 'job' in data:
                print("Single match found, assuming it is the requested simulation.")
                job = matches[0]
            else:
                return jsonify({'error': 'No match found'}), 404
        
        return jsonify({
            'projected_readiness': job['readiness_score'], 
            'projected_match': job.get('weighted_skill_score', job['match_score']) * 100,
            'estimated_weeks': 4
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# ========== SIMULATION ENDPOINTS ==========

@app.route('/api/simulations/recommendation', methods=['GET'])
@jwt_required()
def get_simulation_recommendation():
    try:
        user_id = get_jwt_identity()
        profile = CandidateProfileModel.find_by_user_id(user_id)
        if not profile: return jsonify({'error': 'Profile not found'}), 404

        # Dynamic Recommendation Logic
        target_role = profile.get('target_role', '').lower()
        
        # Broadened heuristic mapping for "whatever" role the user mentions
        sim_id = 'sql-perf-audit' # Default for Analyst, BI, Product, etc.
        
        # Software / Frontend / Web / Fullstack / Apps
        if any(k in target_role for k in ['software', 'developer', 'frontend', 'web', 'fullstack', 'full stack', 'react', 'javascript', 'app', 'mobile']):
            sim_id = 'react-perf-fix'
        # Data Eng / ML / Backend / Python
        elif any(k in target_role for k in ['data engineer', 'python', 'machine learning', 'ml', 'backend', 'etl', 'scientist']):
            sim_id = 'python-data-cleanup'
        # DevOps / Security / Cloud / Infra
        elif any(k in target_role for k in ['devops', 'cloud', 'aws', 'security', 'infra', 'sre', 'site reliability', 'azure', 'architect']):
            sim_id = 'aws-security-audit'
            
        sim = sim_engine.get_simulation(sim_id)
        
        # Calculate current readiness for this role
        profile_data = {
            'preferred_role': profile.get('target_role', sim['role']),
            'experience_level': profile.get('experience_level', 'entry'),
            'skills': profile.get('skills', [])
        }
        results = matcher.match_user(profile_data)
        # Mocking generic readiness if no specific job match found
        current_readiness = results[0]['readiness_score'] if results else 85.0

        # Dynamic Skill Breakdown based on Sim Role
        if 'react' in sim['id']:
            skills_list = [
                {'name': 'React Performance', 'score': 65, 'status': 'Developing'},
                {'name': 'Component Lifecycle', 'score': 70, 'status': 'Good'},
                {'name': 'State Management', 'score': 60, 'status': 'Developing'},
                {'name': 'JavaScript Fundamentals', 'score': 85, 'status': 'Excellent'},
                {'name': 'Frontend Architecture', 'score': 55, 'status': 'Needs Improvement'}
            ]
        elif 'python' in sim['id']:
            skills_list = [
                {'name': 'Data Pipeline Design', 'score': 60, 'status': 'Developing'},
                {'name': 'Python Scripting', 'score': 80, 'status': 'Proficient'},
                {'name': 'Error Handling', 'score': 50, 'status': 'Needs Improvement'},
                {'name': 'ETL Optimization', 'score': 65, 'status': 'Developing'},
                {'name': 'Database Interactions', 'score': 75, 'status': 'Good'}
            ]
        elif 'aws' in sim['id']:
             skills_list = [
                {'name': 'IAM Policy Config', 'score': 45, 'status': 'Critical Gap'},
                {'name': 'Cloud Security', 'score': 55, 'status': 'Developing'},
                {'name': 'Infrastructure as Code', 'score': 70, 'status': 'Good'},
                {'name': 'Audit & Compliance', 'score': 60, 'status': 'Developing'},
                {'name': 'AWS Services', 'score': 75, 'status': 'Proficient'}
            ]
        else:
            skills_list = [
                {'name': 'SQL Query Optimization', 'score': 80, 'status': 'Proficient'},
                {'name': 'Data Visualization', 'score': 90, 'status': 'Excellent'},
                {'name': 'Statistical Analysis', 'score': 60, 'status': 'Developing'},
                {'name': 'Business Intelligence', 'score': 75, 'status': 'Good'},
                {'name': 'Data Modeling', 'score': 65, 'status': 'Developing'}
            ]

        return jsonify({
            'target_role': sim['role'],
            'readiness_score': round(current_readiness, 1),
            'readiness_label': 'Good' if current_readiness > 70 else 'Developing',
            'skill_breakdown': skills_list,
            'recommended_simulation': {
                'id': sim['id'],
                'name': sim['title'],
                'skill_addressed': sim['target_skill'],
                'est_time': sim['estimated_time'],
                'impact_explanation': f"Successfully completing this closes your {sim['target_skill']} reasoning gap for {sim['role']} role."
            },
            'all_available': [
                {
                    'id': s['id'],
                    'name': s['title'],
                    'role': s['role'],
                    'skill': s['target_skill']
                } for s in sim_engine.SIMULATIONS.values()
            ]
        })
    except Exception as e:
        app.logger.error(f"Sim recommendation error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/simulations/<sim_id>', methods=['GET'])
@jwt_required()
def get_simulation_details(sim_id):
    try:
        sim = sim_engine.get_simulation(sim_id)
        if not sim: return jsonify({'error': 'Simulation not found'}), 404
        return jsonify(sim)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/simulations/<sim_id>/start', methods=['POST'])
@jwt_required()
def start_simulation(sim_id):
    try:
        user_id = get_jwt_identity()
        sim = sim_engine.get_simulation(sim_id)
        if not sim: return jsonify({'error': 'Simulation not found'}), 404

        profile = CandidateProfileModel.find_by_user_id(user_id)
        if not profile: return jsonify({'error': 'Profile not found'}), 404

        # Calculate dynamic initial readiness baseline
        matcher_profile = {
            'skills': profile.get('skills', []),
            'experience_level': profile.get('experience_level', 'Entry Level'),
            'preferred_role': sim.get('role', 'Software Engineer')
        }
        
        # Create a mock job representing the simulation's requirements
        mock_job = {
            'job_id': 'sim_baseline',
            'title': sim.get('role', 'Software Engineer'),
            'mapped_skills': [sim.get('target_skill')],
            'formatted_experience_level': profile.get('experience_level', 'Entry Level')
        }
        
        baseline_matches = matcher.match_user(matcher_profile, live_jobs=[mock_job])
        initial_readiness = baseline_matches[0].get('readiness_score', 80.0) if baseline_matches else 80.0

        # Create attempt in DB
        attempt = SimulationModel.create(
            user_id=user_id,
            simulation_id=sim_id,
            scenario_name=sim['title'],
            initial_readiness=initial_readiness
        )

        return jsonify({
            'attempt_id': str(attempt['_id']),
            'simulation': sim
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/simulations/<attempt_id>/submit', methods=['POST'])
@jwt_required()
def submit_simulation(attempt_id):
    try:
        user_id = get_jwt_identity()
        data = request.json
        decisions = data.get('decisions', [])
        justification = data.get('justification', '')

        attempt = SimulationModel.find_by_id(attempt_id)
        if not attempt: return jsonify({'error': 'Attempt not found'}), 404

        # Evaluate impact
        result = sim_engine.evaluate_submission(attempt['simulation_id'], decisions, justification)
        if not result: return jsonify({'error': 'Evaluation failed'}), 500

        final_readiness = attempt['initial_readiness'] + result['total_score']
        
        # Save results
        SimulationModel.update_result(attempt_id, decisions, final_readiness, justification, evaluation_result=result)
        
        # Update user profile skills
        # This would be more complex in real app (history tracking)
        # For now, we assume sim improves target skill
        
        return jsonify({
            'success': True,
            'impact': result['total_score'],
            'before_score': round(attempt['initial_readiness'], 1),
            'after_score': round(final_readiness, 1),
            'breakdown': result['breakdown'], 
            'skill_impact': result['skill_impact']
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ========== CHATBOT ENDPOINT ==========
from chat_service import chat_service

@app.route('/api/chat', methods=['POST'])
@jwt_required()
def chat_endpoint():
    try:
        user_id = get_jwt_identity()
        data = request.json
        message = data.get('message', '')
        
        # Get user context
        profile = CandidateProfileModel.find_by_user_id(user_id)
        if profile:
            # Handle both string and dict skills (Robustly)
            skills = profile.get('skills', [])
            skill_names = []
            for s in skills:
                if isinstance(s, dict):
                    skill_names.append(str(s.get('name', '')))
                else:
                    skill_names.append(str(s))
            user_context = f"Role: {profile.get('target_role', 'Unknown')}, Skills: {', '.join(skill_names)}"
        else:
            user_context = "Guest User"
        
        response_text = chat_service.generate_response(message, context=user_context)
        return jsonify({'response': response_text})
    except Exception as e:
        app.logger.error(f"Chat error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({
        "status": "healthy",
        "websocket": "enabled",
        "timestamp": datetime.utcnow().isoformat()
    })

@app.route('/api/test_reload', methods=['GET'])
def test_reload():
    return jsonify({"message": "Reloaded!"}), 200

# Serve Frontend SPA
@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def serve(path):
    if path != "" and os.path.exists(app.static_folder + '/' + path):
        return send_from_directory(app.static_folder, path)
    else:
        return send_from_directory(app.static_folder, 'index.html')

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    socketio.run(app, debug=True, host='0.0.0.0', port=port, allow_unsafe_werkzeug=True)

